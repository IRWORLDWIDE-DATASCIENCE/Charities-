import fitz  # PyMuPDF
import pdfplumber
import re
import csv
import os
import logging
from transformers import AutoTokenizer, AutoModelForSequenceClassification, pipeline
import torch
from docx import Document

logging.getLogger("pdfminer").setLevel(logging.ERROR)

# Logging config
logging.basicConfig(level=logging.INFO)

# Load FinBERT tokenizer and model
tokenizer = AutoTokenizer.from_pretrained("ProsusAI/finbert")
model = AutoModelForSequenceClassification.from_pretrained("ProsusAI/finbert")

# Load summarizer
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

def extract_project_blocks(text):
    # Use paragraph-level blocks with some filtering logic
    blocks = text.split('\n\n')
    keywords = ["project", "programme", "strategy", "research", "initiative", "campaign", "trial", "funded", "support", "activity"]
    relevant_blocks = [b.strip() for b in blocks if any(k in b.lower() for k in keywords) and len(b.strip()) > 50]
    return relevant_blocks


def extract_text_from_pdf(file_path):
    logging.info("Extracting text from PDF...")
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text


def simple_sent_tokenize(text):
    return re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)


def classify_with_finbert(text_block):
    inputs = tokenizer(text_block, return_tensors="pt", truncation=True, max_length=512)
    with torch.no_grad():
        outputs = model(**inputs)
    predicted_class = torch.argmax(outputs.logits).item()
    labels = ['negative', 'neutral', 'positive']
    return labels[predicted_class]


def classify_context(text_block):
    text_lower = text_block.lower()
    income_keywords = [
        "income", "raised", "revenue", "donated", "donations",
        "gifted", "received", "earned", "awarded", "grants received",
        "funding received", "grant income", "support received", "profit", "help"
    ]
    expenditure_keywords = [
        "expenditure", "spent", "used to", "donated to others",
        "support costs", "charitable activities", "paid",
        "allocated to", "given to", "disbursed", "expenses",
        "cost of", "granted to", "outgoings"
    ]

    if any(keyword in text_lower for keyword in income_keywords):
        return "income"
    elif any(keyword in text_lower for keyword in expenditure_keywords):
        return "expenditure"
    else:
        sentiment = classify_with_finbert(text_block)
        if sentiment == "positive":
            return "income"
        elif sentiment == "negative":
            return "expenditure"
        else:
            return "unknown"


def parse_financial_entries(text):
    logging.info("Parsing financial entries from text...")
    text = re.sub(
        r'£\s*([\d,]+)\s*\n?\.\s*?(\d+)\s*(m|million|k|thousand)?',
        lambda m: f"£{m.group(1)}.{m.group(2)}{m.group(3) or ''}",
        text, flags=re.IGNORECASE
    )

    sentences = simple_sent_tokenize(text)
    entries = []

    for sentence in sentences:
        pattern = r"£\s?([\d,]+(?:\.\d{1,2})?)\s*(million|thousand|m|k)?"
        matches = re.findall(pattern, sentence, re.IGNORECASE)

        for amount_str, scale_str in matches:
            raw_str = f"£{amount_str} {scale_str or ''}".strip()
            try:
                amount = float(amount_str.replace(",", ""))
            except ValueError:
                continue

            scale_str = (scale_str or "").lower()
            if scale_str in ["million", "m"]:
                value = amount * 1_000_000
                scale = "million"
            elif scale_str in ["thousand", "k"]:
                value = amount * 1_000
                scale = "thousand"
            else:
                value = amount
                scale = "million" if amount >= 1_000_000 else "thousand" if amount >= 1_000 else None

            # <-- ADD HERE: Override classification if "annual pay" or "pension" in context
            context_lower = sentence.lower()
            if "annual pay" in context_lower or "pension" in context_lower:
                classification = "expenditure"
            else:
                classification = classify_context(sentence)

            entries.append({
                "raw": raw_str,
                "value": value,
                "scale": scale,
                "context": sentence,
                "type": classification
            })

    return entries


def extract_tables_from_pdf(file_path):
    logging.info("Extracting tables from PDF...")
    table_entries = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    if not row:
                        continue
                    for cell in row:
                        if not cell:
                            continue
                        match = re.search(r'£?\s?[\d,]+(?:\.\d{1,2})?', cell)
                        if match:
                            context = ' '.join(filter(None, row))
                            context_lower = context.lower()
                            # <-- ADD HERE too: Override classification for tables
                            if "annual pay" in context_lower or "pension" in context_lower:
                                classification = "expenditure"
                            else:
                                classification = classify_context(context)
                            try:
                                value = float(match.group().replace("£", "").replace(",", ""))
                            except ValueError:
                                continue
                            table_entries.append({
                                "raw": match.group(),
                                "value": value,
                                "scale": None,
                                "context": context,
                                "type": classification
                            })
    return table_entries

def extract_tables_from_pdf(file_path):
    logging.info("Extracting tables from PDF...")
    table_entries = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                headers = table[0] if table else []
                for row in table[1:]:  # skip header
                    if not row or all(cell is None or cell.strip() == '' for cell in row):
                        continue
                    context = ' | '.join(filter(None, row))
                    for cell in row:
                        if not cell:
                            continue
                        match = re.search(r'£?\s?[\d,]+(?:\.\d{1,2})?', cell)
                        if match:
                            try:
                                value = float(match.group().replace("£", "").replace(",", ""))
                            except ValueError:
                                continue
                            context_lower = context.lower()
                            if "annual pay" in context_lower or "pension" in context_lower:
                                classification = "expenditure"
                            else:
                                classification = classify_context(context)
                            table_entries.append({
                                "raw": match.group(),
                                "value": value,
                                "scale": None,
                                "context": context,
                                "type": classification,
                                "source": "table"
                            })
    return table_entries



def deduplicate_entries(entries):
    logging.info("Deduplicating entries...")
    unique = []
    seen = set()
    for entry in entries:
        sig = (round(entry['value'], 2), entry['type'])
        if sig not in seen:
            unique.append(entry)
            seen.add(sig)
    return unique


def extract_project_sentences(text):
    keywords = ["project", "programme", "program", "initiative", "campaign", "activity", "support", "community", "help", "response"]
    sentences = simple_sent_tokenize(text)
    return [s for s in sentences if any(k in s.lower() for k in keywords)]


def save_to_csv(data, pdf_file_path):
    base_name = os.path.splitext(os.path.basename(pdf_file_path))[0]
    folder = os.path.dirname(pdf_file_path)
    filename = os.path.join(folder, f"{base_name}.csv")

    keys = ["raw", "value", "scale", "context", "type", "source"]
    with open(filename, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=keys)
        writer.writeheader()
        for row in data:
            writer.writerow({
                "raw": row["raw"],
                "value": f"{row['value']:,.2f}",
                "scale": row["scale"] or "units",
                "context": row["context"],
                "type": row["type"],
                "source": row.get("source", "text")
            })

    logging.info(f"✅ Results saved to: {filename}")

def print_summary(entries, projects):
    print("\n📊 Top Financial Entries:\n")
    for i, entry in enumerate(entries, 1):
        print(f"{i}. Amount: {entry['raw']} → £{entry['value']:,.2f}")
        print(f"   Type: {entry['type'].capitalize()} | Scale: {entry['scale'] or 'units'}")
        print(f"   Context: {entry['context'][:200]}{'...' if len(entry['context']) > 200 else ''}\n")

    print("\n📌 Project-related Sentences:\n")
    for s in projects[:10]:  # First 10 for clarity
        print(f" - {s.strip()}")
def generate_financial_summary(entries):
    income_total = sum(e['value'] for e in entries if e['type'] == 'income')
    expenditure_total = sum(e['value'] for e in entries if e['type'] == 'expenditure')

    income_items = sorted([e for e in entries if e['type'] == 'income'], key=lambda x: x['value'], reverse=True)
    expense_items = sorted([e for e in entries if e['type'] == 'expenditure'], key=lambda x: x['value'], reverse=True)

    print("💰 Estimated Total Income: £{:,.2f}".format(income_total))
    print("💸 Estimated Total Expenditure: £{:,.2f}".format(expenditure_total))
    print("\n📈 Top Income Sources:")
    for item in income_items[:5]:
        print(f"  - {item['raw']}: {item['context'][:100]}")

    print("\n📉 Top Expenditures:")
    for item in expense_items[:5]:
        print(f"  - {item['raw']}: {item['context'][:100]}")
        
def save_project_sentences_to_word(sentences, pdf_file_path):
    base_name = os.path.splitext(os.path.basename(pdf_file_path))[0]
    folder = os.path.dirname(pdf_file_path)
    docx_filename = os.path.join(folder, f"{base_name}_project_sentences.docx")

    doc = Document()
    doc.add_heading("Project-Related Sentences", level=1)

    for i, sentence in enumerate(sentences, 1):
        doc.add_paragraph(f"{i}. {sentence.strip()}", style="List Number")

    doc.save(docx_filename)
    logging.info(f"📄 Project sentences saved to: {docx_filename}")

def save_projects_to_word(project_blocks, pdf_file_path):
    base_name = os.path.splitext(os.path.basename(pdf_file_path))[0]
    folder = os.path.dirname(pdf_file_path)
    docx_filename = os.path.join(folder, f"{base_name}_project_summary.docx")

    doc = Document()
    doc.add_heading("Cancer Research UK – 2021/22 Project Summary", level=1)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Project / Focus Area'
    hdr_cells[1].text = 'Details'
    hdr_cells[2].text = 'Funding (if any)'
    hdr_cells[3].text = 'Notes / Outcomes'

    for block in project_blocks:
        row = table.add_row().cells
        row[0].text = block[:50] + "..." if len(block) > 50 else block  # Title
        row[1].text = block  # Full content
        row[2].text = "£..." if "£" in block else "N/A"
        row[3].text = "Ongoing" if "launch" in block.lower() or "plan" in block.lower() else "Reported"

    doc.save(docx_filename)
    logging.info(f"📄 Structured project summary saved to: {docx_filename}")




# ------------------ MAIN ------------------

if __name__ == "__main__":
    pdf_path = r"C:\Users\Naseer.Nijamudin\OneDrive - Islamic Relief Worldwide\Desktop\ALL CHARITIES AUDIT\NON FAITH BASED\Cancer Research Uk\2021 to 2022.pdf"
    # Step 1: Extract text & tables
    text = extract_text_from_pdf(pdf_path)
    text_entries = parse_financial_entries(text)
    table_entries = extract_tables_from_pdf(pdf_path)
    all_entries = deduplicate_entries(text_entries + table_entries)

    # Step 2: Sort top financial entries
    top_entries = sorted(all_entries, key=lambda x: x['value'], reverse=True)
    generate_financial_summary(top_entries)


    # Step 3: Extract project context
    project_sentences = extract_project_sentences(text)

    # Step 4: Display results and save CSV
    print_summary(top_entries, project_sentences)
    save_to_csv(top_entries, pdf_path)
    project_blocks = extract_project_blocks(text)
    save_projects_to_word(project_blocks, pdf_path)


